"""
Enhanced converter using Mermaid.ink API for diagram rendering
No local dependencies needed - uses online API
"""

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Preformatted, Image
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
import base64
import re
import os
import tempfile
import zlib
import json

class MermaidRenderer:
    """Render Mermaid diagrams using mermaid.ink API"""
    
    def __init__(self):
        self.diagram_count = 0
        self.temp_dir = tempfile.mkdtemp()
        print(f"📁 Temp directory for diagrams: {self.temp_dir}")
    
    def encode_mermaid(self, mermaid_code):
        """Encode mermaid code for URL"""
        # Encode to JSON format
        graph_data = {"code": mermaid_code, "mermaid": {"theme": "default"}}
        json_str = json.dumps(graph_data)
        
        # Compress and base64 encode
        compressed = zlib.compress(json_str.encode('utf-8'))
        encoded = base64.urlsafe_b64encode(compressed).decode('ascii')
        
        return encoded
    
    def render_mermaid(self, mermaid_code):
        """Render mermaid code to PNG using mermaid.ink API"""
        try:
            self.diagram_count += 1
            output_file = os.path.join(self.temp_dir, f'diagram_{self.diagram_count}.png')
            
            print(f"  🎨 Rendering diagram {self.diagram_count}...", end='')
            
            # Try simpler kroki.io API first (more reliable)
            try:
                # Use kroki.io API - more reliable
                url = "https://kroki.io/mermaid/png"
                headers = {'Content-Type': 'text/plain'}
                response = requests.post(url, data=mermaid_code.encode('utf-8'), headers=headers, timeout=30)
                
                if response.status_code == 200:
                    with open(output_file, 'wb') as f:
                        f.write(response.content)
                    print(" ✅")
                    return output_file
                else:
                    print(f" ⚠️ (API returned {response.status_code})")
                    return None
                    
            except requests.exceptions.RequestException as e:
                print(f" ❌ (API error: {e})")
                return None
                
        except Exception as e:
            print(f" ❌ (Error: {e})")
            return None

def parse_markdown_with_diagrams(md_file, diagram_renderer):
    """Parse markdown and extract Mermaid diagrams"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    elements = []
    lines = content.split('\n')
    i = 0
    
    print(f"📖 Parsing {len(lines)} lines...")
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Mermaid diagrams
        if line.startswith('```mermaid'):
            diagram_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                diagram_lines.append(lines[i])
                i += 1
            i += 1  # Skip closing ```
            
            mermaid_code = '\n'.join(diagram_lines)
            
            # Try to render the diagram
            img_path = diagram_renderer.render_mermaid(mermaid_code)
            
            if img_path and os.path.exists(img_path):
                elements.append(('diagram', img_path))
            else:
                # Fallback to text description
                title = "Diagram"
                if 'graph' in mermaid_code.lower():
                    if 'Architecture' in mermaid_code or 'Component' in mermaid_code:
                        title = "System Architecture Diagram"
                    elif 'Threat' in mermaid_code:
                        title = "Threat Model Diagram"
                    else:
                        title = "Process Flow Diagram"
                elif 'sequenceDiagram' in mermaid_code:
                    title = "Sequence Diagram"
                elif 'timeline' in mermaid_code:
                    title = "Timeline Diagram"
                
                elements.append(('diagram_placeholder', f'[{title} - See online version or markdown source]'))
            continue
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            elements.append(('heading', level, text))
        
        # Code blocks
        elif line.startswith('```'):
            code_lang = line[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            elements.append(('code', '\n'.join(code_lines)))
        
        # Tables
        elif '|' in line and i+1 < len(lines) and '----|' in lines[i+1]:
            table_lines = [line]
            i += 1  # Skip separator
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            elements.append(('table', table_lines))
            continue
        
        # Lists
        elif line.startswith(('- ', '* ', '+ ', '✅', '❌', '⚠️')) or (line and line[0].isdigit() and '.' in line[:3]):
            elements.append(('list_item', line))
        
        # Horizontal rule
        elif line.startswith('---'):
            elements.append(('hr', None))
        
        # Blockquotes
        elif line.startswith('>'):
            elements.append(('blockquote', line[1:].strip()))
        
        # Regular paragraphs
        else:
            elements.append(('paragraph', line))
        
        i += 1
    
    return elements

def create_pdf_with_diagrams(md_file, output_file):
    """Generate PDF with rendered diagrams"""
    print(f"\n{'='*60}")
    print(f"📄 Creating PDF with diagrams")
    print(f"{'='*60}")
    
    # Initialize diagram renderer
    renderer = MermaidRenderer()
    
    # Create PDF document
    doc = SimpleDocTemplate(
        output_file,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    # Styles
    styles = getSampleStyleSheet()
    
    # Custom styles
    styles.add(ParagraphStyle(
        name='CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1a1a1a'),
        spaceAfter=30,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    ))
    
    styles.add(ParagraphStyle(
        name='CustomHeading1',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#2c3e50'),
        spaceAfter=12,
        spaceBefore=12,
        fontName='Helvetica-Bold'
    ))
    
    styles.add(ParagraphStyle(
        name='CustomHeading2',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#34495e'),
        spaceAfter=10,
        spaceBefore=10,
        fontName='Helvetica-Bold'
    ))
    
    styles.add(ParagraphStyle(
        name='CustomHeading3',
        parent=styles['Heading3'],
        fontSize=12,
        textColor=colors.HexColor('#7f8c8d'),
        spaceAfter=8,
        spaceBefore=8,
        fontName='Helvetica-Bold'
    ))
    
    styles.add(ParagraphStyle(
        name='CustomCode',
        parent=styles['Code'],
        fontSize=8,
        fontName='Courier',
        textColor=colors.HexColor('#2c3e50'),
        backColor=colors.HexColor('#f8f9fa'),
        leftIndent=20,
        rightIndent=20
    ))
    
    # Parse markdown with diagrams
    elements_data = parse_markdown_with_diagrams(md_file, renderer)
    
    # Build PDF elements
    story = []
    
    print(f"\n📝 Building PDF content...")
    
    for item in elements_data:
        if item[0] == 'heading':
            level, text = item[1], item[2]
            
            # Clean markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
            text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
            text = re.sub(r'`(.*?)`', r'<font name="Courier">\1</font>', text)
            
            if level == 1:
                story.append(Paragraph(text, styles['CustomTitle']))
                story.append(Spacer(1, 0.3*inch))
            elif level == 2:
                story.append(Spacer(1, 0.2*inch))
                story.append(Paragraph(text, styles['CustomHeading1']))
            elif level == 3:
                story.append(Paragraph(text, styles['CustomHeading2']))
            else:
                story.append(Paragraph(text, styles['CustomHeading3']))
        
        elif item[0] == 'diagram':
            # Add rendered diagram image
            img_path = item[1]
            try:
                img = Image(img_path, width=5.5*inch, height=4*inch, kind='proportional')
                story.append(img)
                story.append(Spacer(1, 0.25*inch))
            except Exception as e:
                print(f"⚠️ Could not add diagram image: {e}")
                story.append(Paragraph('<i>[Diagram]</i>', styles['Normal']))
        
        elif item[0] == 'diagram_placeholder':
            # Add placeholder text for diagrams that couldn't be rendered
            story.append(Paragraph(f'<i>{item[1]}</i>', styles['Normal']))
            story.append(Spacer(1, 0.15*inch))
        
        elif item[0] == 'paragraph':
            text = item[1]
            text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
            text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
            text = re.sub(r'`(.*?)`', r'<font name="Courier">\1</font>', text)
            
            story.append(Paragraph(text, styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
        
        elif item[0] == 'code':
            code_text = item[1]
            lines = code_text.split('\n')
            if len(lines) > 45:
                code_text = '\n'.join(lines[:45]) + '\n... [code truncated for space]'
            
            story.append(Preformatted(code_text, styles['CustomCode']))
            story.append(Spacer(1, 0.15*inch))
        
        elif item[0] == 'table':
            table_lines = item[1]
            data = []
            for line in table_lines:
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                data.append(cells)
            
            if data:
                t = Table(data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#34495e')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(t)
                story.append(Spacer(1, 0.2*inch))
        
        elif item[0] == 'list_item':
            text = item[1]
            text = re.sub(r'^[\-\*\+]\s+', '• ', text)
            text = re.sub(r'^\d+\.\s+', '• ', text)
            text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
            text = re.sub(r'`(.*?)`', r'<font name="Courier">\1</font>', text)
            
            story.append(Paragraph(text, styles['Normal']))
        
        elif item[0] == 'hr':
            story.append(Spacer(1, 0.1*inch))
            story.append(PageBreak())
        
        elif item[0] == 'blockquote':
            text = item[1]
            story.append(Paragraph(f'<i>"{text}"</i>', styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
    
    # Build PDF
    print(f"🔨 Generating PDF file...")
    doc.build(story)
    print(f"✅ PDF created: {output_file}")
    print(f"   Diagrams rendered: {renderer.diagram_count}")

def create_docx_with_diagrams(md_file, output_file):
    """Generate DOCX with rendered diagrams"""
    print(f"\n{'='*60}")
    print(f"📝 Creating DOCX with diagrams")
    print(f"{'='*60}")
    
    # Initialize diagram renderer
    renderer = MermaidRenderer()
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Parse markdown with diagrams
    blocks = parse_markdown_with_diagrams(md_file, renderer)
    
    # Build document
    print(f"\n📝 Building DOCX content...")
    
    for block in blocks:
        if block[0] == 'heading':
            level, text = block[1], block[2]
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            
            if level == 1:
                p = doc.add_heading(text, level=0)
            elif level == 2:
                doc.add_heading(text, level=1)
            elif level == 3:
                doc.add_heading(text, level=2)
            else:
                doc.add_heading(text, level=3)
        
        elif block[0] == 'diagram':
            # Add rendered diagram image
            img_path = block[1]
            try:
                doc.add_picture(img_path, width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()  # Spacing
            except Exception as e:
                print(f"⚠️ Could not add diagram to DOCX: {e}")
                p = doc.add_paragraph('[Diagram]')
                for run in p.runs:
                    run.italic = True
        
        elif block[0] == 'diagram_placeholder':
            p = doc.add_paragraph(block[1])
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
        
        elif block[0] == 'paragraph':
            text = block[1]
            p = doc.add_paragraph()
            
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                else:
                    p.add_run(part)
        
        elif block[0] == 'code':
            code_text = block[1]
            lines = code_text.split('\n')
            if len(lines) > 40:
                code_text = '\n'.join(lines[:40]) + '\n... [code truncated]'
            
            p = doc.add_paragraph(code_text)
            p.style = 'No Spacing'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
        
        elif block[0] == 'table':
            table_lines = block[1]
            data = []
            for line in table_lines:
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                data.append(cells)
            
            if data:
                table = doc.add_table(rows=len(data), cols=len(data[0]))
                table.style = 'Light Grid Accent 1'
                
                for i, row_data in enumerate(data):
                    cells = table.rows[i].cells
                    for j, cell_text in enumerate(row_data):
                        cell_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                        cell_text = re.sub(r'`(.*?)`', r'\1', cell_text)
                        cells[j].text = cell_text
                        
                        if i == 0:
                            for paragraph in cells[j].paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                
                doc.add_paragraph()
        
        elif block[0] == 'list_item':
            text = block[1]
            text = re.sub(r'^[\-\*\+✅❌⚠️]\s+', '', text)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            
            doc.add_paragraph(text, style='List Bullet')
        
        elif block[0] == 'hr':
            doc.add_page_break()
        
        elif block[0] == 'blockquote':
            text = block[1]
            p = doc.add_paragraph(f'"{text}"')
            for run in p.runs:
                run.italic = True
    
    # Save document
    print(f"🔨 Generating DOCX file...")
    doc.save(output_file)
    print(f"✅ DOCX created: {output_file}")
    print(f"   Diagrams rendered: {renderer.diagram_count}")

if __name__ == '__main__':
    print("\n" + "="*60)
    print("📊 WHITEPAPER CONVERTER WITH MERMAID DIAGRAMS")
    print("   Using Kroki.io API for diagram rendering")
    print("="*60)
    
    # Generate both formats with diagrams
    try:
        create_pdf_with_diagrams('WHITEPAPER_COMPLETE.md', 'WHITEPAPER_COMPLETE.pdf')
        create_docx_with_diagrams('WHITEPAPER_COMPLETE.md', 'WHITEPAPER_COMPLETE.docx')
        
        print("\n" + "="*60)
        print("✅ CONVERSION COMPLETE!")
        print("="*60)
        print("\n📋 Files generated:")
        print("   • WHITEPAPER_COMPLETE.pdf (with diagrams)")
        print("   • WHITEPAPER_COMPLETE.docx (with diagrams)")
        print("\n💡 Note: If some diagrams didn't render, check your")
        print("   internet connection (requires API access)")
        print("="*60 + "\n")
        
    except Exception as e:
        print(f"\n❌ Error during conversion: {e}")
        import traceback
        traceback.print_exc()
