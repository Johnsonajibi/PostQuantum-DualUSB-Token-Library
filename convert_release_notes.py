"""
Convert RELEASE_NOTES_v0.1.1.md to PDF and DOCX formats
"""

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Preformatted
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def parse_release_notes(md_file):
    """Parse release notes markdown"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    elements = []
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            elements.append(('heading', level, text))
        
        # Code blocks
        elif line.startswith('```'):
            code_lang = line[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            elements.append(('code', '\n'.join(code_lines)))
        
        # Lists
        elif line.startswith(('- ', '* ', '+ ')):
            elements.append(('list_item', line))
        
        # Regular paragraphs
        else:
            elements.append(('paragraph', line))
        
        i += 1
    
    return elements

def create_pdf(md_file, output_file):
    """Generate PDF from release notes"""
    print(f"\n📄 Creating PDF: {output_file}")
    
    doc = SimpleDocTemplate(
        output_file,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    # Styles
    styles = getSampleStyleSheet()
    
    styles.add(ParagraphStyle(
        name='CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        textColor=colors.HexColor('#1a1a1a'),
        spaceAfter=20,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    ))
    
    styles.add(ParagraphStyle(
        name='CustomH2',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=colors.HexColor('#2c3e50'),
        spaceAfter=12,
        spaceBefore=15,
        fontName='Helvetica-Bold'
    ))
    
    styles.add(ParagraphStyle(
        name='CustomH3',
        parent=styles['Heading2'],
        fontSize=13,
        textColor=colors.HexColor('#34495e'),
        spaceAfter=10,
        spaceBefore=10,
        fontName='Helvetica-Bold'
    ))
    
    styles.add(ParagraphStyle(
        name='CustomCode',
        parent=styles['Code'],
        fontSize=9,
        fontName='Courier',
        textColor=colors.HexColor('#2c3e50'),
        backColor=colors.HexColor('#f8f9fa'),
        leftIndent=20,
        rightIndent=20,
        spaceAfter=10
    ))
    
    # Parse markdown
    elements_data = parse_release_notes(md_file)
    
    # Build PDF
    story = []
    
    for item in elements_data:
        if item[0] == 'heading':
            level, text = item[1], item[2]
            
            # Clean markdown
            text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
            text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
            text = re.sub(r'`(.*?)`', r'<font name="Courier">\1</font>', text)
            
            if level == 1:
                story.append(Paragraph(text, styles['CustomTitle']))
                story.append(Spacer(1, 0.2*inch))
            elif level == 2:
                story.append(Paragraph(text, styles['CustomH2']))
            else:
                story.append(Paragraph(text, styles['CustomH3']))
        
        elif item[0] == 'paragraph':
            text = item[1]
            text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
            text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
            text = re.sub(r'`(.*?)`', r'<font name="Courier">\1</font>', text)
            
            # Handle links
            text = re.sub(r'\[(.*?)\]\((.*?)\)', r'<link href="\2">\1</link>', text)
            
            story.append(Paragraph(text, styles['Normal']))
            story.append(Spacer(1, 0.08*inch))
        
        elif item[0] == 'code':
            code_text = item[1]
            story.append(Preformatted(code_text, styles['CustomCode']))
            story.append(Spacer(1, 0.1*inch))
        
        elif item[0] == 'list_item':
            text = item[1]
            text = re.sub(r'^[\-\*\+]\s+', '• ', text)
            text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
            text = re.sub(r'`(.*?)`', r'<font name="Courier">\1</font>', text)
            
            story.append(Paragraph(text, styles['Normal']))
    
    # Build PDF
    doc.build(story)
    print(f"✅ PDF created: {output_file}")

def create_docx(md_file, output_file):
    """Generate DOCX from release notes"""
    print(f"\n📝 Creating DOCX: {output_file}")
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Parse markdown
    blocks = parse_release_notes(md_file)
    
    # Build document
    for block in blocks:
        if block[0] == 'heading':
            level, text = block[1], block[2]
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            
            if level == 1:
                p = doc.add_heading(text, level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                doc.add_heading(text, level=1)
            else:
                doc.add_heading(text, level=2)
        
        elif block[0] == 'paragraph':
            text = block[1]
            p = doc.add_paragraph()
            
            # Simple bold/italic handling
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                elif part.startswith('[') and '](' in part:
                    # Handle links
                    match = re.match(r'\[(.*?)\]\((.*?)\)', part)
                    if match:
                        run = p.add_run(match.group(1))
                        run.font.color.rgb = RGBColor(0, 0, 255)
                        run.underline = True
                else:
                    p.add_run(part)
        
        elif block[0] == 'code':
            code_text = block[1]
            p = doc.add_paragraph(code_text)
            p.style = 'No Spacing'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
        
        elif block[0] == 'list_item':
            text = block[1]
            text = re.sub(r'^[\-\*\+]\s+', '', text)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            
            doc.add_paragraph(text, style='List Bullet')
    
    # Save document
    doc.save(output_file)
    print(f"✅ DOCX created: {output_file}")

if __name__ == '__main__':
    print("="*60)
    print("📋 RELEASE NOTES CONVERTER")
    print("="*60)
    
    # Generate both formats
    create_pdf('RELEASE_NOTES_v0.1.1.md', 'RELEASE_NOTES_v0.1.1.pdf')
    create_docx('RELEASE_NOTES_v0.1.1.md', 'RELEASE_NOTES_v0.1.1.docx')
    
    print("\n" + "="*60)
    print("✅ CONVERSION COMPLETE!")
    print("="*60)
    print("\n📋 Files generated:")
    print("   • RELEASE_NOTES_v0.1.1.pdf")
    print("   • RELEASE_NOTES_v0.1.1.docx")
    print("="*60 + "\n")
