"""
Alternative converter using Python libraries (markdown2, python-docx, reportlab)
Install: pip install markdown2 python-docx reportlab
"""

def convert_with_python_libs():
    """Convert using pure Python libraries"""
    from pathlib import Path
    import subprocess
    import sys
    
    # Check if required packages are installed
    required_packages = {
        'markdown2': 'markdown2',
        'docx': 'python-docx',
        'reportlab': 'reportlab'
    }
    
    missing_packages = []
    for module, package in required_packages.items():
        try:
            __import__(module)
        except ImportError:
            missing_packages.append(package)
    
    if missing_packages:
        print("Installing required packages...")
        print(f"pip install {' '.join(missing_packages)}")
        response = input("Install now? (y/n): ")
        if response.lower() == 'y':
            subprocess.check_call([sys.executable, '-m', 'pip', 'install'] + missing_packages)
            print("✓ Packages installed. Please run the script again.")
            return
        else:
            print("Aborted. Please install packages manually:")
            print(f"  pip install {' '.join(missing_packages)}")
            return
    
    # Now do the conversion
    import markdown2
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re
    
    whitepaper_path = Path(__file__).parent / "WHITEPAPER.md"
    
    if not whitepaper_path.exists():
        print(f"Error: {whitepaper_path} not found!")
        return
    
    print("Reading markdown file...")
    with open(whitepaper_path, 'r', encoding='utf-8') as f:
        markdown_content = f.read()
    
    # Convert to DOCX
    print("\nConverting to Microsoft Word (DOCX)...")
    try:
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Parse markdown and add to document
        lines = markdown_content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip horizontal rules
            if line.startswith('---'):
                i += 1
                continue
            
            # Main title (# )
            if line.startswith('# ') and not line.startswith('## '):
                text = line.replace('# ', '')
                heading = doc.add_heading(text, level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Heading level 2 (## )
            elif line.startswith('## '):
                text = line.replace('## ', '')
                doc.add_heading(text, level=2)
            
            # Heading level 3 (### )
            elif line.startswith('### '):
                text = line.replace('### ', '')
                doc.add_heading(text, level=3)
            
            # Bold text (**text**)
            elif '**' in line:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part:
                        p.add_run(part)
            
            # Code blocks (```)
            elif line.startswith('```'):
                # Collect code block
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                
                # Add code block with monospace font
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'Normal'
                run = p.runs[0] if p.runs else p.add_run()
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            
            # Bullet points (- )
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:]
                doc.add_paragraph(text, style='List Bullet')
            
            # Numbered lists (1. )
            elif re.match(r'^\d+\. ', line):
                text = re.sub(r'^\d+\. ', '', line)
                doc.add_paragraph(text, style='List Number')
            
            # Tables (| )
            elif line.startswith('|'):
                # Collect table rows
                table_rows = [line]
                i += 1
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_rows.append(lines[i].strip())
                    i += 1
                i -= 1
                
                # Parse table
                rows = []
                for row in table_rows:
                    if '---' not in row:  # Skip separator row
                        cells = [cell.strip() for cell in row.split('|')[1:-1]]
                        rows.append(cells)
                
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_data in enumerate(row_data):
                            table.rows[row_idx].cells[col_idx].text = cell_data
            
            # Normal paragraph
            elif line:
                doc.add_paragraph(line)
            
            i += 1
        
        # Save DOCX
        docx_output = whitepaper_path.with_suffix('.docx')
        doc.save(docx_output)
        print(f"✓ DOCX created: {docx_output}")
        
    except Exception as e:
        print(f"✗ Error creating DOCX: {e}")
    
    print("\n" + "="*60)
    print("Note: For PDF conversion, use convert_whitepaper.py with pandoc")
    print("Or use Microsoft Word to save the DOCX as PDF")
    print("="*60)

if __name__ == "__main__":
    convert_with_python_libs()
