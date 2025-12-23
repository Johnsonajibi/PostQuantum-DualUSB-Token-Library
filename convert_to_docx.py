"""
Convert WHITEPAPER_COMPLETE.md to Microsoft Word DOCX
Uses python-docx for professional document generation
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def parse_markdown_simple(md_file):
    """Parse markdown into simple structure"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    blocks = []
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip mermaid diagrams
        if line.strip().startswith('```mermaid'):
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
            i += 1
            blocks.append(('note', '[Diagram - see markdown source]'))
            continue
        
        # Code blocks
        if line.strip().startswith('```'):
            lang = line.strip()[3:]
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            blocks.append(('code', '\n'.join(code_lines)))
            i += 1
            continue
        
        # Headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            blocks.append(('heading', level, text))
        
        # Tables
        elif '|' in line and i+1 < len(lines) and '---|' in lines[i+1]:
            table_lines = [line]
            i += 1  # Skip separator
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            blocks.append(('table', table_lines))
            continue
        
        # Lists
        elif line.strip().startswith(('- ', '* ', '+ ', '✅', '❌', '⚠️')):
            blocks.append(('list', line.strip()))
        
        # Horizontal rules
        elif line.strip() == '---':
            blocks.append(('hr', None))
        
        # Blockquotes
        elif line.strip().startswith('>'):
            blocks.append(('quote', line.strip()[1:].strip()))
        
        # Regular text
        elif line.strip():
            blocks.append(('text', line.strip()))
        
        i += 1
    
    return blocks

def create_docx(md_file, output_file):
    """Generate DOCX from markdown"""
    print(f"Converting {md_file} to {output_file}...")
    
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Define custom styles
    styles = doc.styles
    
    # Title style
    if 'Custom Title' not in styles:
        title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 0, 0)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
    
    # Parse markdown
    blocks = parse_markdown_simple(md_file)
    
    # Build document
    for block in blocks:
        if block[0] == 'heading':
            level, text = block[1], block[2]
            
            # Clean markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            
            if level == 1:
                p = doc.add_heading(text, level=0)
                p.style = 'Custom Title' if 'Custom Title' in styles else 'Title'
            elif level == 2:
                doc.add_heading(text, level=1)
            elif level == 3:
                doc.add_heading(text, level=2)
            else:
                doc.add_heading(text, level=3)
        
        elif block[0] == 'text':
            text = block[1]
            
            # Format inline markdown
            p = doc.add_paragraph()
            
            # Simple bold/italic handling
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                else:
                    p.add_run(part)
        
        elif block[0] == 'code':
            code_text = block[1]
            # Limit code size
            lines = code_text.split('\n')
            if len(lines) > 40:
                code_text = '\n'.join(lines[:40]) + '\n... [truncated]'
            
            p = doc.add_paragraph(code_text)
            p.style = 'No Spacing'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
        
        elif block[0] == 'table':
            table_lines = block[1]
            # Parse table data
            data = []
            for line in table_lines:
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                data.append(cells)
            
            if data:
                # Create table
                table = doc.add_table(rows=len(data), cols=len(data[0]))
                table.style = 'Light Grid Accent 1'
                
                # Fill table
                for i, row_data in enumerate(data):
                    cells = table.rows[i].cells
                    for j, cell_text in enumerate(row_data):
                        # Clean markdown
                        cell_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                        cell_text = re.sub(r'`(.*?)`', r'\1', cell_text)
                        cells[j].text = cell_text
                        
                        # Bold header row
                        if i == 0:
                            for paragraph in cells[j].paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                
                doc.add_paragraph()  # Spacing after table
        
        elif block[0] == 'list':
            text = block[1]
            # Remove list markers
            text = re.sub(r'^[\-\*\+✅❌⚠️]\s+', '', text)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            
            doc.add_paragraph(text, style='List Bullet')
        
        elif block[0] == 'hr':
            doc.add_page_break()
        
        elif block[0] == 'quote':
            text = block[1]
            p = doc.add_paragraph(f'"{text}"')
            for run in p.runs:
                run.italic = True
        
        elif block[0] == 'note':
            p = doc.add_paragraph(block[1])
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save document
    print("Saving DOCX...")
    doc.save(output_file)
    print(f"✓ DOCX created: {output_file}")

if __name__ == '__main__':
    create_docx('WHITEPAPER_COMPLETE.md', 'WHITEPAPER_COMPLETE.docx')
